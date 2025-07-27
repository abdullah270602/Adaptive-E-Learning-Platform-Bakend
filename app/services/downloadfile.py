import io
import json
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def create_pdf(mcqs: list) -> io.BytesIO:
    """
    Generate PDF from MCQ list with debugging
    """
    print(f"DEBUG: MCQs received: {len(mcqs) if mcqs else 0}")
    print(f"DEBUG: MCQs type: {type(mcqs)}")
    
    if mcqs:
        print(f"DEBUG: First MCQ: {mcqs[0]}")
    
    if not mcqs:
        raise ValueError("No MCQs provided for PDF generation")
    
    buffer = io.BytesIO()
    
    try:
        # Simple document setup
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        story.append(Paragraph("Multiple Choice Questions", styles['Title']))
        story.append(Spacer(1, 20))
        
        # Add each MCQ with better error handling
        for i, mcq in enumerate(mcqs, 1):
            print(f"DEBUG: Processing MCQ {i}: {mcq}")
            
            # Question - handle different possible keys
            question = mcq.get('question') or mcq.get('Question') or f"Question {i} not found"
            story.append(Paragraph(f"{i}. {question}", styles['Heading2']))
            story.append(Spacer(1, 10))
            
            # Options - handle different possible formats
            options = mcq.get('options') or mcq.get('Options') or []
            if isinstance(options, str):
                # If options is a string, try to parse or split it
                try:
                    options = json.loads(options)
                except:
                    options = [options]  # Treat as single option
            
            for option in options:
                story.append(Paragraph(f"   {option}", styles['Normal']))
            
            story.append(Spacer(1, 10))
            
            # Answer
            # answer = mcq.get('correct_answer') or mcq.get('answer') or mcq.get('Answer') or "No answer provided"
            # story.append(Paragraph(f"Answer: {answer}", styles['Normal']))
            story.append(Spacer(1, 20))
        
        print(f"DEBUG: Story items created: {len(story)}")
        
        # Build the PDF
        doc.build(story)
        
        # Check buffer size
        buffer_size = buffer.tell()
        print(f"DEBUG: Buffer size after build: {buffer_size}")
        
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"DEBUG: Error in PDF creation: {str(e)}")
        buffer.close()
        raise RuntimeError(f"PDF generation failed: {str(e)}")

def create_docx(mcqs: list) -> io.BytesIO:
    """
    Generate DOCX from MCQ list with debugging
    """
    from docx import Document
    
    print(f"DEBUG: DOCX MCQs received: {len(mcqs) if mcqs else 0}")
    
    if not mcqs:
        raise ValueError("No MCQs provided for DOCX generation")
    
    buffer = io.BytesIO()
    
    try:
        doc = Document()
        
        # Add title
        doc.add_heading('Multiple Choice Questions', 0)
        
        # Add each MCQ
        for i, mcq in enumerate(mcqs, 1):
            # Question
            question = mcq.get('question') or mcq.get('Question') or f"Question {i} not found"
            doc.add_heading(f"{i}. {question}", level=2)
            
            # Options
            options = mcq.get('options') or mcq.get('Options') or []
            if isinstance(options, str):
                try:
                    options = json.loads(options)
                except:
                    options = [options]
            
            for option in options:
                doc.add_paragraph(f"   {option}")
            
            # Answer
            # answer = mcq.get('correct_answer') or mcq.get('answer') or mcq.get('Answer') or "No answer provided"
            # doc.add_paragraph(f"Answer: {answer}")
            doc.add_paragraph()  # Add spacing
        
        # Save to buffer
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"DEBUG: Error in DOCX creation: {str(e)}")
        buffer.close()
        raise RuntimeError(f"DOCX generation failed: {str(e)}")